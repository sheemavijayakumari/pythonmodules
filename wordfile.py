elif type == "qr_students_image":
            id=db_admin.institution_id            
            ins_name="""Select institution_name from institution where institution_id='{institutionid}'""".format(institutionid=id)
            ins_name= cursor_details(ins_name)
            ins_name_strings = [' '.join(d.values()) if isinstance(d, dict) else str(d) for d in ins_name]
            ins_name_string = ' '.join(ins_name_strings)
            ins_address="""Select address from institution where institution_id='{institutionid}'""".format(institutionid=id)
            ins_address= cursor_details(ins_address)
            ins_address_strings = [' '.join(d.values()) if isinstance(d, dict) else str(d) for d in ins_address]
            ins_address_string = ' '.join(ins_address_strings)
            doc_file_name = "QR_STUDENT_IMAGE.doc"
            if bus_id=='All buses':
                bus_id=None
            if class_id=='All Classes':
                class_id=None
            if bus_id!=None and class_id!=None:
                query = """Select A.student_name,A.class_number,A.qr_filename,C.bus_number from "{student_table}" as A
                            left join "{bus_allocation}" as B ON A.id = ANY(B.student_id)
                            left join "{bus_table}" as C on C.id = B.bus_id
                            left join "{class_table}" as D on D.class_number = A.class_number
                            where C.id = '{t_bus_id}' and D.id = '{t_class_id}'""".format(student_table=db_admin.institution_id+"_student",class_table=db_admin.institution_id+"_class",bus_allocation=db_admin.institution_id+"_bus_allocation",bus_table=db_admin.institution_id+"_bus",t_bus_id=bus_id,t_class_id=class_id)
                student_image = cursor_details(query)
            if bus_id!=None and class_id==None:
                query = """Select A.student_name,A.class_number,A.qr_filename,C.bus_number from "{student_table}" as A
                            left join "{bus_allocation}" as B ON A.id = ANY(B.student_id)
                            left join "{bus_table}" as C on C.id = B.bus_id
                            left join "{class_table}" as D on D.class_number = A.class_number
                            where C.id = '{t_bus_id}'""".format(student_table=db_admin.institution_id+"_student",bus_allocation=db_admin.institution_id+"_bus_allocation",class_table=db_admin.institution_id+"_class",bus_table=db_admin.institution_id+"_bus",t_bus_id=bus_id)
                student_image = cursor_details(query)
            if bus_id==None and class_id!=None:
                query = """Select A.student_name,A.class_number,A.qr_filename,C.bus_number from "{student_table}" as A
                            left join "{bus_allocation}" as B ON A.id = ANY(B.student_id)
                            left join "{bus_table}" as C on C.id = B.bus_id
                            left join "{class_table}" as D on D.class_number = A.class_number                           
                            where A.class_number = '{t_class_id}'""".format(student_table=db_admin.institution_id+"_student",bus_allocation=db_admin.institution_id+"_bus_allocation",class_table=db_admin.institution_id+"_class",bus_table=db_admin.institution_id+"_bus",t_class_id=class_id)
                student_image = cursor_details(query)
            if bus_id==None and class_id==None:
                query = """Select A.student_name,A.class_number,A.qr_filename,C.bus_number from "{student_table}" as A
                            left join "{bus_allocation}" as B ON A.id = ANY(B.student_id)
                            left join "{bus_table}" as C on C.id = B.bus_id
                            left join "{class_table}" as D on D.class_number = A.class_number                            
                            """.format(student_table=db_admin.institution_id+"_student",bus_allocation=db_admin.institution_id+"_bus_allocation",class_table=db_admin.institution_id+"_class",bus_table=db_admin.institution_id+"_bus")
                student_image = cursor_details(query)
            import boto3
            from io import BytesIO
            from docx import Document
            from docx.shared import Inches
            from docx.oxml import OxmlElement, ns
            from docx.shared import Pt
            from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
            aws_access_key_id = 'AKIA26OFCRYRXI6K7NOZ'
            aws_secret_access_key = '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'
            aws_bucket_name = 'gotoz'
            s3 = boto3.client('s3', aws_access_key_id=aws_access_key_id, aws_secret_access_key=aws_secret_access_key)
            doc = Document()
            section = doc.sections[0]
            section.page_width = Inches(8.27) 
            section.page_height = Inches(11.69)
            image_count = 0
            row = len(student_image)            
            rows = row
            cols = 3
            table = doc.add_table(rows=rows, cols=cols)
            for aws_key in student_image:
                if aws_key['qr_filename'] is not None:
                    qr_filename = aws_key['qr_filename'].split('?')[0]
                    response = s3.get_object(Bucket=aws_bucket_name, Key=qr_filename)
                    image_data = response['Body'].read()
                    image_stream = BytesIO(image_data)
                    row_index = image_count // 3
                    col_index = image_count % 3
                    cell = table.cell(row_index, col_index)
                    cell.width = Inches(2)
                    cell.height = Inches(2)
                    for border_side in ('top', 'bottom', 'start', 'end'):
                        # Define the namespace for the XML tags
                        side_border = OxmlElement('w:{}'.format(border_side))
                        side_border.set(ns.qn('w:val'), 'single')
                        side_border.set(ns.qn('w:sz'), '4')
                        side_border.set(ns.qn('w:space'), '0')
                        side_border.set(ns.qn('w:color'), 'auto')
                        cell._element.tcPr.append(side_border)
                        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                    cell.paragraphs[0].add_run().add_picture(image_stream, width=Inches(2),height=Inches(2))
                    cell.paragraphs[0].add_run(f"{ins_name_string}").bold = True 
                    cell.paragraphs[0].runs[-1].font.size = Pt(9)
                    cell.paragraphs[0].add_run(f"\nName: {aws_key['student_name']}\nClass: {aws_key['class_number']}\nBus Number: {aws_key['bus_number']}")
                    image_count += 1
            doc.save(f"{file_path}/QR_STUDENT_IMAGE.doc")
            # import subprocess            
            # subprocess.run(["unoconv", "-f", "pdf", f"{file_path}/QR_STUDENT_IMAGE.docx"])
            # subprocess.run(["xdg-open", f"{file_path}/QR_STUDENT_IMAGE.pdf"])
            return {"file_name":f"{db_admin.institution_id}/{doc_file_name}","records" : [],"records1":[],"records2":[],"ins_name":ins_name_string,"ins_address":ins_address_string}